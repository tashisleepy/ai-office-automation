#!/usr/bin/env python3
"""
python-docx Template Generator
Design a master report template once → auto-fill with any data forever.
Feed it forensic data, client research, or AI analysis — get consulting-grade reports.

Usage:
  python3 scripts/docx-template.py
  python3 scripts/docx-template.py --output my-report.docx

Dependencies: pip3 install python-docx
"""

import argparse
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# ── BRAND CONFIG ──
BRAND = {
    "primary": RGBColor(0x1E, 0x27, 0x61),
    "secondary": RGBColor(0xCA, 0xDC, 0xFC),
    "text_dark": RGBColor(0x33, 0x33, 0x33),
    "text_muted": RGBColor(0x99, 0x99, 0x99),
    "font_title": "Arial",
    "font_body": "Calibri",
}


def set_cell_shading(cell, color_hex):
    """Apply background shading to a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear'
    })
    shading.append(shading_elem)


def build_report(data, output="output-report.docx"):
    """Build a full report from a data dictionary."""
    doc = Document()

    # ── Page margins ──
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # ── Default font ──
    style = doc.styles['Normal']
    font = style.font
    font.name = BRAND["font_body"]
    font.size = Pt(11)
    font.color.rgb = BRAND["text_dark"]

    # ── Heading styles ──
    for level, size in [(1, 18), (2, 14), (3, 12)]:
        h_style = doc.styles[f'Heading {level}']
        h_style.font.name = BRAND["font_title"]
        h_style.font.size = Pt(size)
        h_style.font.color.rgb = BRAND["primary"]
        h_style.font.bold = True

    # ── Title ──
    title_para = doc.add_heading(data["title"], level=1)

    # ── Divider ──
    divider = doc.add_paragraph()
    divider.paragraph_format.space_after = Pt(12)

    # ── Overview ──
    if "overview" in data:
        p = doc.add_paragraph()
        for segment in data["overview"]:
            run = p.add_run(segment["text"])
            if segment.get("bold"):
                run.bold = True
            run.font.size = Pt(11)
            run.font.name = BRAND["font_body"]

    # ── Sections ──
    for section in data.get("sections", []):
        doc.add_heading(section["heading"], level=2)

        # Paragraphs
        for para_data in section.get("paragraphs", []):
            p = doc.add_paragraph()
            for segment in para_data:
                run = p.add_run(segment["text"])
                if segment.get("bold"):
                    run.bold = True
                run.font.size = Pt(11)
                run.font.name = BRAND["font_body"]

        # Bullet points
        for bullet in section.get("bullets", []):
            p = doc.add_paragraph(bullet, style='List Bullet')
            for run in p.runs:
                run.font.size = Pt(11)
                run.font.name = BRAND["font_body"]

        # Table
        if "table" in section:
            tbl_data = section["table"]
            table = doc.add_table(
                rows=1 + len(tbl_data["rows"]),
                cols=len(tbl_data["headers"])
            )
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = 'Table Grid'

            # Header row
            for j, header in enumerate(tbl_data["headers"]):
                cell = table.rows[0].cells[j]
                cell.text = header
                set_cell_shading(cell, "1E2761")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        run.font.size = Pt(10)
                        run.font.name = BRAND["font_body"]

            # Data rows
            for i, row in enumerate(tbl_data["rows"]):
                for j, val in enumerate(row):
                    cell = table.rows[i + 1].cells[j]
                    cell.text = str(val)
                    if i % 2 == 1:
                        set_cell_shading(cell, "F5F7FA")
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10)
                            run.font.name = BRAND["font_body"]

            doc.add_paragraph()  # spacing after table

    # ── Contact ──
    if "contact" in data:
        doc.add_heading("Contact", level=2)
        for key, val in data["contact"].items():
            p = doc.add_paragraph()
            run_key = p.add_run(f"{key}: ")
            run_key.bold = True
            run_key.font.color.rgb = BRAND["primary"]
            run_key.font.name = BRAND["font_body"]
            run_val = p.add_run(val)
            run_val.font.name = BRAND["font_body"]

    doc.save(output)
    print(f"Report created: {output}")


# ── EXAMPLE DATA ──
if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--output", "-o", default="output-report.docx")
    args = parser.parse_args()

    example_data = {
        "title": "Your Company \u2014 Company Overview",
        "overview": [
            {"text": "Your Company", "bold": True},
            {"text": " is an enterprise AI solutions company that has deployed $10M+ in production systems"},
            {"text": " for major conglomerates and global media companies."},
        ],
        "sections": [
            {
                "heading": "Services",
                "table": {
                    "headers": ["Service", "Description", "Delivery"],
                    "rows": [
                        ["AI Consulting", "End-to-end AI strategy and deployment", "4-12 weeks"],
                        ["Competitive Intelligence", "Forensic teardowns with funding reality", "1-2 weeks"],
                        ["Product Development", "Production-ready AI systems at scale", "8-24 weeks"],
                    ]
                }
            },
            {
                "heading": "Key Clients",
                "bullets": [
                    "Media Corp A \u2014 AI video production pipeline",
                    "Media Corp B \u2014 Content automation systems",
                    "Sports Franchise \u2014 AI operations",
                    "Retail Group \u2014 Enterprise AI systems",
                    "Energy Corp \u2014 AI solutions deployment",
                    "Consumer Brand \u2014 AI integration",
                ]
            },
        ],
        "contact": {
            "Email": "hello@example.com",
            "Phone": "+91 XXXXXXXXXX",
            "Location": "Your City, Your Country",
        }
    }

    build_report(example_data, args.output)
